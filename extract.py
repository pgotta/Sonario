"""
extract.py — turn any supported file into plain text.

Formats: .txt .md .markdown .rtf  (text)
         .docx                     (python-docx)
         .pdf                      (pdfplumber text layer; OCR fallback if empty)
         .png .jpg .jpeg .tif .bmp (OCR via pytesseract)

OCR requires Tesseract installed on the machine; if it's missing we skip those
files gracefully rather than crashing the whole run.
"""

import os
import hashlib
import glob
import shutil

SUPPORTED = {
    ".txt", ".md", ".markdown", ".text", ".log",
    ".rtf", ".docx", ".pdf",
    ".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp",
}

IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"}

# Tesseract is often installed (e.g. by winget) but not on the current process's
# PATH, so pytesseract can't find it and OCR fails even though the .exe exists.
# Locate it directly in the standard install locations and point pytesseract at it.
_TESSERACT_READY = None  # None=unchecked, True=found+configured, False=not found


def _configure_tesseract():
    """Find tesseract.exe and tell pytesseract where it is. Returns True if OCR
    is usable. Cached after the first call."""
    global _TESSERACT_READY
    if _TESSERACT_READY is not None:
        return _TESSERACT_READY
    try:
        import pytesseract
    except ImportError:
        _TESSERACT_READY = False
        return False

    # 1) Already on PATH? (the normal case once PATH has refreshed)
    exe = shutil.which("tesseract")
    if exe:
        pytesseract.pytesseract.tesseract_cmd = exe
        _TESSERACT_READY = True
        return True

    # 2) Search the standard Windows install locations, including winget's path.
    candidates = []
    for base in [
        os.environ.get("ProgramFiles", r"C:\Program Files"),
        os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)"),
        os.environ.get("LOCALAPPDATA", ""),
    ]:
        if not base:
            continue
        candidates.append(os.path.join(base, "Tesseract-OCR", "tesseract.exe"))
    # winget installs under ...\WinGet\Packages\UB-Mannheim.TesseractOCR_*\...
    local = os.environ.get("LOCALAPPDATA", "")
    if local:
        candidates += glob.glob(os.path.join(
            local, "Microsoft", "WinGet", "Packages",
            "UB-Mannheim.TesseractOCR*", "**", "tesseract.exe"), recursive=True)
    # also any winget link shim
    if local:
        candidates.append(os.path.join(local, "Microsoft", "WinGet", "Links",
                                       "tesseract.exe"))

    for c in candidates:
        if c and os.path.isfile(c):
            pytesseract.pytesseract.tesseract_cmd = c
            _TESSERACT_READY = True
            return True

    _TESSERACT_READY = False
    return False


def file_hash(path):
    """Stable hash of path + mtime + size — cache key that invalidates on edit."""
    st = os.stat(path)
    key = f"{os.path.abspath(path)}|{st.st_mtime_ns}|{st.st_size}"
    return hashlib.sha256(key.encode()).hexdigest()[:16]


def walk_folder(root):
    """Recursively yield supported file paths under root (skips hidden dirs)."""
    found = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = [d for d in dirnames if not d.startswith(".")]
        for fn in filenames:
            ext = os.path.splitext(fn)[1].lower()
            if ext in SUPPORTED:
                found.append(os.path.join(dirpath, fn))
    return sorted(found)


def extract_text(path):
    """Return (text, note). text may be '' if nothing extractable. note explains why."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in {".txt", ".md", ".markdown", ".text", ".log"}:
            return _read_plain(path), ""
        if ext == ".rtf":
            return _read_rtf(path), ""
        if ext == ".docx":
            return _read_docx(path), ""
        if ext == ".pdf":
            return _read_pdf(path)
        if ext in IMAGE_EXTS:
            return _ocr_image(path)
    except Exception as e:
        return "", f"extract error: {e}"
    return "", "unsupported"


def _read_plain(path):
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    with open(path, "rb") as f:
        return f.read().decode("utf-8", errors="replace")


def _read_rtf(path):
    from striprtf.striprtf import rtf_to_text
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return rtf_to_text(f.read())


def _read_docx(path):
    import docx
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_pdf(path):
    import pdfplumber
    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            text_parts.append(txt)
    text = "\n".join(text_parts).strip()
    if len(text) >= 30:
        return text, ""
    # No usable text layer — likely scanned. Try OCR per page.
    ocr_text, note = _ocr_pdf(path)
    if ocr_text:
        return ocr_text, "ocr"
    return text, note or "pdf had no extractable text"


def _ocr_pdf(path):
    try:
        import pytesseract
        from pdf2image import convert_from_path
    except ImportError:
        return "", "OCR Python libs missing — run setup.bat"
    if not _configure_tesseract():
        return "", ("Tesseract OCR is installed but could not be located. Try "
                    "restarting Sonario (run.bat) in a fresh window, or reboot so "
                    "Windows picks up the new PATH. If it persists, reinstall via "
                    "setup.bat and click YES on the admin prompt.")
    try:
        pages = convert_from_path(path, dpi=200)
    except Exception as e:
        return "", ("Poppler not found (needed for scanned PDFs) — run "
                    "setup.bat to install it")
    out = []
    for img in pages:
        try:
            out.append(pytesseract.image_to_string(img))
        except Exception as e:
            return "", (f"OCR failed: {e}. Tesseract may need a restart of Sonario "
                        "or a reboot to be picked up.")
    return "\n".join(out).strip(), ""


def _ocr_image(path):
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "", "OCR Python libs missing — run setup.bat"
    if not _configure_tesseract():
        return "", ("Tesseract OCR is installed but could not be located. Try "
                    "restarting Sonario (run.bat) in a fresh window, or reboot so "
                    "Windows picks up the new PATH. If it persists, reinstall via "
                    "setup.bat and click YES on the admin prompt.")
    try:
        text = pytesseract.image_to_string(Image.open(path)).strip()
        return text, ("ocr" if text else "no text found in image")
    except Exception as e:
        return "", (f"OCR failed: {e}. Tesseract may need a restart of Sonario "
                    "or a reboot to be picked up.")
